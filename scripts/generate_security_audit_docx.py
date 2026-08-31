#!/usr/bin/env python3
"""Generate a professional Word document from the TimeFlow security audit report."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "SECURITY_AUDIT_REPORT.md"
OUT_PATH = ROOT / "TimeFlow_Security_Audit_Report.docx"
OUT_PATH_VERSIONED = ROOT / "TimeFlow_Security_Audit_Report_v1.3.docx"
IMPROVED_BG = "E8F5E9"

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0xC4, 0x39, 0x2B)  # severity critical accent
MUTED = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ROW_ALT = "F5F7FA"
HEADER_BG = "1B2A4A"
CRITICAL_BG = "FDECEC"
HIGH_BG = "FFF4E5"
MEDIUM_BG = "FFFBEA"
LOW_BG = "F0F7F4"


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def set_cell_shading(cell, hex_color: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "CBD2DC")
        borders.append(element)
    tbl_pr.append(borders)


def shade_severity_row(row, severity: str):
    sev = severity.lower()
    if "critical" in sev:
        color = CRITICAL_BG
    elif "high" in sev:
        color = HIGH_BG
    elif "medium" in sev:
        color = MEDIUM_BG
    elif "low" in sev or "info" in sev:
        color = LOW_BG
    else:
        return
    for cell in row.cells:
        set_cell_shading(cell, color)


def add_formatted_paragraph(doc, text: str, style=None, space_after=8):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    add_inline_runs(p, text)
    return p


def add_inline_runs(paragraph, text: str):
    # Split on **bold**, `code`, and plain text
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True, size=11, color=NAVY)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=9.5, color=RGBColor(0x33, 0x33, 0x33))
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=11)


def add_heading_custom(doc, text: str, level: int):
    heading = doc.add_heading(level=level)
    heading.clear()
    run = heading.add_run(text)
    if level == 1:
        set_run_font(run, name="Calibri", size=16, bold=True, color=NAVY)
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(10)
    elif level == 2:
        set_run_font(run, name="Calibri", size=13, bold=True, color=NAVY)
        heading.paragraph_format.space_before = Pt(14)
        heading.paragraph_format.space_after = Pt(8)
    else:
        set_run_font(run, name="Calibri", size=11.5, bold=True, color=NAVY)
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(6)
    return heading


def add_table(doc, headers, rows, shade_by_severity_col=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=WHITE)
        set_cell_shading(hdr_cells[i], HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, value in enumerate(row_data):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            # strip markdown bold/code for table cells for cleanliness
            clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
            clean = re.sub(r"`([^`]+)`", r"\1", clean)
            run = p.add_run(clean)
            set_run_font(run, size=9.5)
        if shade_by_severity_col is not None and shade_by_severity_col < len(row_data):
            shade_severity_row(table.rows[r_idx + 1], row_data[shade_by_severity_col])
        elif r_idx % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, ROW_ALT)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def parse_md_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (headers, rows, next_idx)."""
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append(lines[i].strip())
        i += 1
    if len(rows) < 2:
        return None, None, start_idx

    def split_row(line):
        parts = [c.strip() for c in line.strip("|").split("|")]
        return parts

    headers = split_row(rows[0])
    # skip separator
    body = []
    for row in rows[2:]:
        cols = split_row(row)
        # pad/truncate
        if len(cols) < len(headers):
            cols += [""] * (len(headers) - len(cols))
        body.append(cols[: len(headers)])
    return headers, body, i


def build_cover(doc: Document):
    # Classification banner
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = banner.add_run("CONFIDENTIAL — INTERNAL USE ONLY")
    set_run_font(run, size=10, bold=True, color=ACCENT)

    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SECURITY AUDIT REPORT")
    set_run_font(run, name="Calibri", size=28, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("TimeFlow Platform")
    set_run_font(run, name="Calibri", size=20, bold=True, color=NAVY)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Web Application • Electron Desktop • Supabase • CI/CD")
    set_run_font(run, size=12, color=MUTED)

    for _ in range(2):
        doc.add_paragraph()

    meta_items = [
        ("Engagement Type", "Comparative re-verification (report-only); owner confirmed key rotation"),
        ("Report Date", "31 August 2026"),
        ("Prior Reports", "v1.0 / v1.1 — 24 Aug 2026; v1.2 — 28 Aug 2026"),
        ("Overall Posture", "Improved (Moderate)"),
        ("Version", "1.3"),
        ("Latest roll-up", "Resolved / Unresolved (+ detail: Fixed, Partial, Mitigated, …)"),
        ("Classification", "Confidential — Internal Use Only"),
        ("Methodology", "OWASP ASVS L2 (tailored) • OWASP Top 10 / API Top 10 • Supabase model • GitHub Actions • Selected CIS • Comparative status tables"),
    ]

    table = doc.add_table(rows=len(meta_items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for i, (k, v) in enumerate(meta_items):
        table.rows[i].cells[0].text = ""
        table.rows[i].cells[1].text = ""
        p0 = table.rows[i].cells[0].paragraphs[0]
        r0 = p0.add_run(k)
        set_run_font(r0, size=10, bold=True, color=WHITE)
        set_cell_shading(table.rows[i].cells[0], HEADER_BG)
        p1 = table.rows[i].cells[1].paragraphs[0]
        r1 = p1.add_run(v)
        set_run_font(r1, size=10)
        if k == "Overall Posture":
            set_cell_shading(table.rows[i].cells[1], IMPROVED_BG)

    for _ in range(3):
        doc.add_paragraph()

    footer_note = doc.add_paragraph()
    footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_note.add_run(
        "This document contains security-sensitive findings. Distribute only to authorized personnel.\n"
        "No formal ASVS/CIS certification is claimed."
    )
    set_run_font(run, size=9, italic=True, color=MUTED)

    doc.add_page_break()


def setup_header_footer(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("TimeFlow Security Audit Report  |  Confidential")
    set_run_font(run, size=8, color=MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Page ")
    set_run_font(run, size=8, color=MUTED)
    add_page_number(fp)
    run2 = fp.add_run("  •  Comparative re-verification  •  31 August 2026  •  v1.3")
    set_run_font(run2, size=8, color=MUTED)


def convert_markdown(doc: Document, md_text: str):
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    skip_title = True  # cover already has title

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip leading H1 and first meta table (cover page duplicates)
        if skip_title:
            if stripped.startswith("# "):
                i += 1
                continue
            if stripped.startswith("|"):
                _, _, i = parse_md_table(lines, i)
                # skip following ---
                while i < len(lines) and lines[i].strip() in ("", "---"):
                    i += 1
                skip_title = False
                continue
            if stripped == "---" or stripped == "":
                i += 1
                continue
            skip_title = False

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
                lang = stripped[3:].strip()
                if lang == "mermaid":
                    # skip mermaid block; replace with prose once closed
                    i += 1
                    while i < len(lines) and not lines[i].strip().startswith("```"):
                        i += 1
                    i += 1
                    add_formatted_paragraph(
                        doc,
                        "Architecture (simplified): TimeFlow web SPA and Electron desktop clients communicate "
                        "directly with Supabase (Auth, database, storage) and with the external host "
                        "timeflowstorage.mechlintech.com. GitHub Actions builds and deploys the web app to a "
                        "self-hosted Docker/nginx runtime.",
                    )
                    in_code = False
                    continue
            else:
                # flush code block
                code = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.left_indent = Cm(0.3)
                run = p.add_run(code)
                set_run_font(run, name="Consolas", size=8.5, color=RGBColor(0x22, 0x22, 0x22))
                # light background via shading on paragraph is awkward; keep monospace block
                in_code = False
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("|"):
            headers, body, next_i = parse_md_table(lines, i)
            if headers and body is not None:
                sev_col = None
                for idx, h in enumerate(headers):
                    if h.lower() in ("severity", "posture", "rating"):
                        sev_col = idx
                        break
                add_table(doc, headers, body, shade_by_severity_col=sev_col)
                i = next_i
                continue

        if stripped.startswith("### "):
            add_heading_custom(doc, stripped[4:], level=2)
            i += 1
            continue

        if stripped.startswith("## "):
            add_heading_custom(doc, stripped[3:], level=1)
            i += 1
            continue

        if stripped.startswith("# "):
            add_heading_custom(doc, stripped[2:], level=1)
            i += 1
            continue

        if re.match(r"^[-*] ", stripped):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            add_inline_runs(p, text)
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)
            add_inline_runs(p, text)
            i += 1
            continue

        if stripped == "":
            i += 1
            continue

        add_formatted_paragraph(doc, stripped)
        i += 1


def main():
    md_text = MD_PATH.read_text(encoding="utf-8")
    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    build_cover(doc)
    setup_header_footer(doc)
    convert_markdown(doc, md_text)

    # Closing classification
    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end.add_run("— End of Report —")
    set_run_font(run, size=10, italic=True, color=MUTED)

    end2 = doc.add_paragraph()
    end2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end2.add_run("CONFIDENTIAL — INTERNAL USE ONLY")
    set_run_font(run, size=9, bold=True, color=ACCENT)

    doc.save(OUT_PATH)
    doc.save(OUT_PATH_VERSIONED)
    print(f"Wrote {OUT_PATH}")
    print(f"Wrote {OUT_PATH_VERSIONED}")


if __name__ == "__main__":
    main()
